import os
import re
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_hex):
    """Set the background color of a table cell."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """Set padding/margins for a table cell."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('w:top', top), ('w:bottom', bottom), ('w:left', left), ('w:right', right)]:
        node = OxmlElement(m)
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def create_document():
    doc = Document()
    
    # ── Page Setup ────────────────────────────────────────────────────────────
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    
    # Configure styles
    styles = doc.styles
    
    # Normal Style (Body Text)
    style_normal = styles['Normal']
    style_normal.font.name = 'Calibri'
    style_normal.font.size = Pt(11)
    style_normal.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    style_normal.paragraph_format.line_spacing = 1.15
    style_normal.paragraph_format.space_after = Pt(6)
    
    # Heading 1 Style
    style_h1 = styles['Heading 1']
    style_h1.font.name = 'Calibri Light'
    style_h1.font.size = Pt(18)
    style_h1.font.bold = True
    style_h1.font.color.rgb = RGBColor(0x0a, 0x5a, 0x1b) # Dark Green accent
    style_h1.paragraph_format.space_before = Pt(12)
    style_h1.paragraph_format.space_after = Pt(6)
    
    # Heading 2 Style
    style_h2 = styles['Heading 2']
    style_h2.font.name = 'Calibri Light'
    style_h2.font.size = Pt(14)
    style_h2.font.bold = True
    style_h2.font.color.rgb = RGBColor(0xf5, 0x9e, 0x0b) # Orange accent
    style_h2.paragraph_format.space_before = Pt(10)
    style_h2.paragraph_format.space_after = Pt(4)

    # Heading 3 Style
    style_h3 = styles['Heading 3']
    style_h3.font.name = 'Calibri'
    style_h3.font.size = Pt(12)
    style_h3.font.bold = True
    style_h3.font.italic = True
    style_h3.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    style_h3.paragraph_format.space_before = Pt(6)
    style_h3.paragraph_format.space_after = Pt(2)

    # ── Title Page ────────────────────────────────────────────────────────────
    # Add title spacing
    for _ in range(5):
        doc.add_paragraph()
        
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = p_title.add_run("FARMALERT")
    run_title.font.name = 'Calibri'
    run_title.font.size = Pt(36)
    run_title.font.bold = True
    run_title.font.color.rgb = RGBColor(0x0a, 0x5a, 0x1b)
    
    p_tagline = doc.add_paragraph()
    p_tagline.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_tagline = p_tagline.add_run("Hyper-local Climate Risk Forecasts, AI Adaptation Plans, and Satellite-Driven Crop Insights for Smallholder Farmers")
    run_tagline.font.name = 'Calibri'
    run_tagline.font.size = Pt(14)
    run_tagline.font.italic = True
    run_tagline.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    
    for _ in range(8):
        doc.add_paragraph()
        
    p_meta = doc.add_paragraph()
    p_meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_meta = p_meta.add_run("Project Documentation & Technical Architecture Report\n")
    run_meta.font.bold = True
    run_meta.font.size = Pt(12)
    
    run_meta_details = p_meta.add_run(
        "Version: 1.0.0\n"
        "Date: June 10, 2026\n"
        "Author: Senior Technical Engineering & Product Architecture Team\n"
        "Target: Presentation / Technical Submission Ready"
    )
    run_meta_details.font.size = Pt(10)
    run_meta_details.font.color.rgb = RGBColor(0x77, 0x77, 0x77)
    
    doc.add_page_break()

    # ── Read Markdown Content ──────────────────────────────────────────────────
    md_path = "PROJECT_DOCUMENTATION.md"
    if not os.path.exists(md_path):
        print(f"Error: {md_path} not found.")
        return
        
    with open(md_path, 'r', encoding='utf-8') as f:
        content = f.read()

    # Split by major sections (headings starting with # )
    sections = re.split(r'\n(?=# \d+)', content)
    
    for sec in sections:
        sec = sec.strip()
        if not sec:
            continue
            
        lines = sec.split('\n')
        # The first line is the heading (# Heading)
        h_line = lines[0]
        h_text = re.sub(r'^#\s+', '', h_line).strip()
        
        # Add heading
        h = doc.add_heading(h_text, level=1)
        h.paragraph_format.keep_with_next = True
        
        # Parse the rest of the lines
        in_table = False
        table_headers = []
        table_rows = []
        in_list = False
        in_code = False
        code_block = []

        i = 1
        while i < len(lines):
            line = lines[i]
            stripped = line.strip()
            
            # Code block detection
            if stripped.startswith('```'):
                if not in_code:
                    in_code = True
                    code_block = []
                else:
                    in_code = False
                    # Write code block as single block with specific styling
                    p_code = doc.add_paragraph()
                    p_code.paragraph_format.left_indent = Inches(0.25)
                    p_code.paragraph_format.right_indent = Inches(0.25)
                    p_code.paragraph_format.space_before = Pt(4)
                    p_code.paragraph_format.space_after = Pt(4)
                    
                    # Add background shading
                    pBdr = OxmlElement('w:pBdr')
                    left_border = OxmlElement('w:left')
                    left_border.set(qn('w:val'), 'single')
                    left_border.set(qn('w:sz'), '24') # 3pt width
                    left_border.set(qn('w:space'), '12')
                    left_border.set(qn('w:color'), '555555')
                    pBdr.append(left_border)
                    p_code._p.get_or_add_pPr().append(pBdr)
                    
                    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F8F9FA"/>')
                    p_code._p.get_or_add_pPr().append(shd)

                    run_code = p_code.add_run('\n'.join(code_block))
                    run_code.font.name = 'Consolas'
                    run_code.font.size = Pt(9.5)
                    run_code.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
                i += 1
                continue

            if in_code:
                code_block.append(line)
                i += 1
                continue

            # Section Separators (horizontal rules)
            if stripped == '---':
                doc.add_page_break()
                i += 1
                continue

            # Heading 2 detection
            if stripped.startswith('## '):
                h2_text = re.sub(r'^##\s+', '', stripped).strip()
                h2 = doc.add_heading(h2_text, level=2)
                h2.paragraph_format.keep_with_next = True
                i += 1
                continue
                
            # Heading 3 detection
            if stripped.startswith('### '):
                h3_text = re.sub(r'^###\s+', '', stripped).strip()
                h3 = doc.add_heading(h3_text, level=3)
                h3.paragraph_format.keep_with_next = True
                i += 1
                continue
                
            # Table detection
            if stripped.startswith('|'):
                # Check if this is the separator row
                if '---' in stripped:
                    i += 1
                    continue
                
                row_cells = [cell.strip() for cell in stripped.split('|')[1:-1]]
                if not in_table:
                    in_table = True
                    table_headers = row_cells
                    table_rows = []
                else:
                    table_rows.append(row_cells)
                
                # Look ahead to see if the table continues
                next_is_table = False
                if i + 1 < len(lines):
                    next_stripped = lines[i+1].strip()
                    if next_stripped.startswith('|'):
                        next_is_table = True
                
                if not next_is_table:
                    # Render the accumulated table
                    in_table = False
                    col_count = len(table_headers)
                    table = doc.add_table(rows=1, cols=col_count)
                    table.style = 'Light Shading Accent 1' # Word built-in clean table style
                    
                    # Style headers
                    hdr_cells = table.rows[0].cells
                    for col_idx, text in enumerate(table_headers):
                        hdr_cells[col_idx].text = text
                        set_cell_background(hdr_cells[col_idx], "0A5A1B")
                        set_cell_margins(hdr_cells[col_idx], top=120, bottom=120, left=150, right=150)
                        
                        # Style text inside cell
                        for paragraph in hdr_cells[col_idx].paragraphs:
                            for run in paragraph.runs:
                                run.font.bold = True
                                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                                
                    # Fill data rows
                    for r_idx, row_data in enumerate(table_rows):
                        row = table.add_row()
                        row_cells = row.cells
                        bg_color = "F9FBF9" if r_idx % 2 == 0 else "FFFFFF"
                        for col_idx, text in enumerate(row_data):
                            if col_idx < len(row_cells):
                                row_cells[col_idx].text = text
                                set_cell_background(row_cells[col_idx], bg_color)
                                set_cell_margins(row_cells[col_idx], top=100, bottom=100, left=150, right=150)
                                for paragraph in row_cells[col_idx].paragraphs:
                                    paragraph.runs[0].font.size = Pt(10)
                                    paragraph.runs[0].font.color.rgb = RGBColor(0x33, 0x33, 0x33)
                                    
                    # Add paragraph space after table
                    doc.add_paragraph()
                i += 1
                continue

            # List detection
            if stripped.startswith('* ') or stripped.startswith('- ') or (stripped.startswith('1. ') or stripped.startswith('2. ') or stripped.startswith('3. ') or stripped.startswith('4. ') or stripped.startswith('5. ') or stripped.startswith('6. ')):
                is_ordered = stripped[0].isdigit()
                list_text = re.sub(r'^(\*|-|\d+\.)\s+', '', stripped).strip()
                
                # Format bold markers inside lists
                p_item = doc.add_paragraph(style='List Bullet' if not is_ordered else 'List Number')
                p_item.paragraph_format.space_after = Pt(3)
                
                parts = re.split(r'(\*\*.*?\*\*)', list_text)
                for part in parts:
                    if part.startswith('**') and part.endswith('**'):
                        bold_text = part[2:-2]
                        run = p_item.add_run(bold_text)
                        run.font.bold = True
                    else:
                        p_item.add_run(part)
                i += 1
                continue

            # Plain paragraph text
            if stripped:
                p = doc.add_paragraph()
                # Process bold formatting **text**
                parts = re.split(r'(\*\*.*?\*\*)', line)
                for part in parts:
                    if part.startswith('**') and part.endswith('**'):
                        bold_text = part[2:-2]
                        run = p.add_run(bold_text)
                        run.font.bold = True
                    else:
                        p.add_run(part)
            i += 1

    doc.save("PROJECT_DOCUMENTATION.docx")
    print("Successfully compiled PROJECT_DOCUMENTATION.docx")

if __name__ == "__main__":
    create_document()
